"""
core/docx_builder.py
Génère synthese.docx depuis la liste des EcoleCommission en session.
Un tableau par élève, code couleur niveau, mention bilangue, collège rouge si pas SMAC.
"""
from __future__ import annotations
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from core.models import Eleve, Niveau, TypeAide, EcoleCommission

# ─── Couleurs ─────────────────────────────────────────────────────────────────
COULEURS_NIVEAU = {
    Niveau.TRES_BON:   ("1A5C1A", "FFFFFF"),  # vert foncé / blanc
    Niveau.BON:        ("4CAF50", "FFFFFF"),  # vert clair / blanc
    Niveau.MOYEN:      ("FFC107", "000000"),  # jaune / noir
    Niveau.FRAGILE:    ("F44336", "FFFFFF"),  # rouge / blanc
    Niveau.INCONNU:    ("CCCCCC", "000000"),  # gris / noir
}

COULEUR_BILANGUE    = "87CEEB"   # bleu ciel
COULEUR_COLLEGE     = "F44336"   # rouge
COULEUR_ENTETE      = "FFFFFF"   # blanc (au lieu de noir)
COULEUR_ENTETE_TEXT = "000000"   # texte noir
COULEUR_LABEL_BG    = "E0E0E0"   # gris clair (au lieu de anthracite)
COULEUR_LABEL_FG    = "000000"   # texte noir
COULEUR_SMAC_BG     = "E8F5E9"   # vert très pâle pour collège SMAC
COULEUR_SMAC_FG     = "1A5C1A"

# ─── Helpers XML ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Applique une couleur de fond à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell, color: str = "000000", size: str = "4") -> None:
    """Bordures sur toutes les faces."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def _keep_paragraph_together(para) -> None:
    """Empêche la coupure d'un paragraphe entre deux pages."""
    pPr = para._p.get_or_add_pPr()
    # keepLines : empêche la coupure du paragraphe lui-même
    keepLines = OxmlElement("w:keepLines")
    pPr.append(keepLines)
    # keepNext : garde le paragraphe avec le suivant
    keepNext = OxmlElement("w:keepNext")
    pPr.append(keepNext)

def _keep_table_together(table) -> None:
    """Empêche la coupure du tableau entre deux pages."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _keep_paragraph_together(para)

def _para(cell, text: str, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Vide la cellule et écrit un paragraphe."""
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    
    # ← AJOUTER CECI : empêche la coupure du paragraphe
    _keep_paragraph_together(para)
    
    return para

# ─── Construction d'un tableau élève ──────────────────────────────────────────

_LABELS = [
    ("Sexe",             lambda e: "Garçon" if e.sexe.value == "G" else "Fille"),
    ("Niveau",           lambda e: e.niveau.value),
    ("Aide",             lambda e: e.aide.value if e.aide.value else ""),
    ("Bilangue",         lambda e: "OUI" if e.est_bilangue else ""),
    ("Section Basket",   lambda e: "OUI" if e.est_basket else ""),
    ("Pénible",          lambda e: "OUI" if e.est_penible else ""),
    ("À séparer de",     lambda e: ", ".join(e.a_separer_de) if e.a_separer_de else ""),
    ("À regrouper avec", lambda e: ", ".join(e.a_regrouper_avec) if e.a_regrouper_avec else ""),
    ("Collège",          lambda e: e.college),
]

def _build_eleve_table(doc: Document, eleve: Eleve) -> None:
    """Ajoute un tableau complet pour un élève."""
    
    # ── En-tête : NOM Prénom [BILANGUE] ──────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=1)
    header_table.style = "Table Grid"
    hcell = header_table.rows[0].cells[0]
    
    # Fond blanc, bordures noires
    _set_cell_bg(hcell, COULEUR_ENTETE)
    _set_cell_borders(hcell, "000000", "8")  # bordures plus épaisses
    
    # Nom en gras noir
    hcell.paragraphs[0].clear()
    hpara = hcell.paragraphs[0]
    hpara.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run_nom = hpara.add_run(f"{eleve.nom.upper()} {eleve.prenom}")
    run_nom.bold = True
    run_nom.font.size = Pt(12)
    run_nom.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    # 🇬🇧 Drapeau Union Jack pour les élèves bilangues
    if eleve.est_bilangue:
        run_bil = hpara.add_run("  BILANGUE")
        run_bil.bold = True
        run_bil.font.size = Pt(10)
        run_bil.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    # 🏀 Ballon de basket pour les élèves en section basket  
    if eleve.est_basket:
        run_bask = hpara.add_run("  🏀 BASKET")
        run_bask.bold = True
        run_bask.font.size = Pt(10)
        # Violet
        run_bask.font.color.rgb = RGBColor(0x9C, 0x27, 0xB0)
    # ── Corps : tableau 2 colonnes ─────────────────────────────────────────────
    body_table = doc.add_table(rows=len(_LABELS), cols=2)
    body_table.style = "Table Grid"
    
    # Largeurs colonnes
    for row in body_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)
    
    for i, (label, getter) in enumerate(_LABELS):
        valeur = getter(eleve)
        lcell = body_table.rows[i].cells[0]
        vcell = body_table.rows[i].cells[1]
        
        _set_cell_borders(lcell, "000000")
        _set_cell_borders(vcell, "000000")
        _set_cell_bg(lcell, COULEUR_LABEL_BG)
        _set_cell_bg(vcell, "FFFFFF")  # blanc par défaut, écrasé si besoin
        
        # Label
        _para(lcell, label, bold=True, color=COULEUR_LABEL_FG, size=9)
        
        # Valeur avec traitement spécial selon le champ
        if label == "Niveau":
            bg, fg = COULEURS_NIVEAU.get(eleve.niveau, ("CCCCCC", "000000"))
            _set_cell_bg(vcell, bg)
            _para(vcell, valeur or "?", bold=True, color=fg, size=10)
        
        elif label == "Collège":
            is_smac = valeur.upper() in ("SMAC", "") or "MARIE" in valeur.upper() or "ERNEST" in valeur.upper()
            if is_smac:
                _set_cell_bg(vcell, COULEUR_SMAC_BG)
                _para(vcell, "SMAC", color=COULEUR_SMAC_FG, size=10)
            else:
                _set_cell_bg(vcell, COULEUR_COLLEGE)
                _para(vcell, valeur, bold=True, color="FFFFFF", size=10)
        
        elif label == "Bilangue" and valeur == "OUI":
            _set_cell_bg(vcell, "BBDEFB")  # Bleu ciel très clair
            _para(vcell, "OUI", bold=True, color="1565C0", size=10)  # Bleu foncé
        
        elif label == "Section Basket" and valeur == "OUI":
            _set_cell_bg(vcell, "E1BEE7")  # Violet clair
            _para(vcell, "OUI", bold=True, color="7B1FA2", size=10)  # Violet foncé

        elif label == "Pénible" and valeur == "OUI":
            _set_cell_bg(vcell, "FF5722")
            _para(vcell, "OUI", bold=True, color="FFFFFF", size=10)
        
        else:
            _set_cell_bg(vcell, "FFFFFF")
            _para(vcell, valeur, size=10)
    
    # ── Observation : ligne fusionnée sous le tableau ────────────────────────
    if eleve.observation:
        obs_table = doc.add_table(rows=1, cols=1)
        obs_table.style = "Table Grid"
        ocell = obs_table.rows[0].cells[0]
        ocell.width = Cm(16)  # largeur totale
        
        _set_cell_borders(ocell, "000000")
        _set_cell_bg(ocell, "F5F5F5")  # gris très clair
        
        # Label "Observation" en gras
        obs_para = ocell.paragraphs[0]
        run_label = obs_para.add_run("Observation : ")
        run_label.bold = True
        run_label.font.size = Pt(9)
        
        run_obs = obs_para.add_run(eleve.observation)
        run_obs.font.size = Pt(9)
    
    # Empêcher la coupure des tableaux
    _keep_table_together(header_table)
    _keep_table_together(body_table)
    if eleve.observation:
        _keep_table_together(obs_table)
    
    # Espace après chaque élève
    doc.add_paragraph()

# ─── Numérotation des pages ───────────────────────────────────────────────────

def _add_page_numbers(doc: Document) -> None:
    """Ajoute la numérotation des pages en footer."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    for section in doc.sections:
        footer = section.footer
        # Créer un paragraphe centré dans le footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le numéro de page
        run = para.add_run("Page ")
        
        # Champ PAGE pour numéro automatique
        rPr = OxmlElement("w:rPr")
        rPr.set(qn("w:noProof"), "1")
        
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        
        # Ajouter au run
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        para.add_run(" sur ")
        
        # Champ NUMPAGES pour total
        run2 = para.add_run("")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = " NUMPAGES "
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")
        run2._r.append(fldChar3)
        run2._r.append(instrText2)
        run2._r.append(fldChar4)

# ─── API publique ──────────────────────────────────────────────────────────────

def build_synthese_docx(commissions: list[EcoleCommission]) -> bytes:
    """
    Génère synthese.docx depuis toutes les commissions importées.
    Retourne les bytes du fichier prêt à télécharger.
    """
    doc = Document()
    
    # ── Styles globaux ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    
    # Marges de page réduites pour maximiser l'espace
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)
    
    # ── Titre général ─────────────────────────────────────────────────────────
    titre = doc.add_heading("Synthèse Commissions CM2 — 6e", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    total = sum(len(c.eleves) for c in commissions)
    sous_titre = doc.add_paragraph(f"{len(commissions)} école(s) · {total} élèves")
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ── Une section par école ──────────────────────────────────────────────────
    for i, comm in enumerate(commissions):
        # Titre école
        h = doc.add_heading(f"{comm.nom_ecole} — {comm.commune}", level=2)
        p = doc.add_paragraph(f"{len(comm.eleves)} élèves")
        
        # Un tableau par élève
        for eleve in comm.eleves:
            _build_eleve_table(doc, eleve)
        
        # Saut de page entre écoles (pas après la dernière)
        if i < len(commissions) - 1:
            doc.add_page_break()
    
    # ── Numérotation des pages ────────────────────────────────────────────────
    _add_page_numbers(doc)
    
    # ── Export bytes ──────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()